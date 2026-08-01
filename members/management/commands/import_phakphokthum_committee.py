<<<<<<< HEAD
from pathlib import Path
import re
=======
import csv
import hashlib
import re
from pathlib import Path
>>>>>>> 1d670fd (refactor)

from django.conf import settings
from django.core.management.base import BaseCommand, CommandError
from django.db import transaction
from docx import Document

<<<<<<< HEAD
from members.models import Member
from members.name_utils import romanize_nepali_name
=======
from members.models import Member, MembershipRecord
from members.services import get_or_create_category, get_or_create_unit, normalize_phone, resolve_person_identity, sync_legacy_member
>>>>>>> 1d670fd (refactor)


NEPALI_DIGITS = str.maketrans("०१२३४५६७८९", "0123456789")
DEFAULT_FILE = Path("data") / "Phakphokthum Rural Commeetee.docx"
DEFAULT_UNIT = "Phakphokthum Rural Municipality"


def clean_text(value):
    return " ".join(str(value or "").replace("\xa0", " ").split()).strip()


def nepali_number_to_int(value):
    text = clean_text(value).translate(NEPALI_DIGITS)
    match = re.search(r"\d+", text)
    return int(match.group()) if match else 0


<<<<<<< HEAD
def clean_phone(value):
    text = clean_text(value).translate(NEPALI_DIGITS)
    return re.sub(r"[^0-9+]", "", text)


class Command(BaseCommand):
    help = (
        "Import the 11-person Phakphokthum rural committee DOCX into the "
        "public members.Member registry."
    )

    def add_arguments(self, parser):
        parser.add_argument(
            "--file",
            default=str(DEFAULT_FILE),
            help="DOCX path relative to manage.py, or an absolute path.",
        )
        parser.add_argument(
            "--membership-type",
            choices=["general", "life"],
            default="general",
            help=(
                "The source document does not state membership type. "
                "The importer defaults to general."
            ),
        )
        parser.add_argument(
            "--unit-name",
            default=DEFAULT_UNIT,
            help="Registry/unit name shown in the member-directory filter.",
        )
        parser.add_argument(
            "--show-phone",
            action="store_true",
            help="Show imported phone numbers publicly. Default: hidden.",
        )
        parser.add_argument(
            "--replace-unit",
            action="store_true",
            help="Delete current records in this exact registry before importing.",
        )
=======
class Command(BaseCommand):
    help = "Safely import the Phakphokthum rural committee DOCX without deleting existing members."

    def add_arguments(self, parser):
        parser.add_argument("--file", default=str(DEFAULT_FILE))
        parser.add_argument("--membership-type", choices=["general", "life"], default="general")
        parser.add_argument("--unit-name", default=DEFAULT_UNIT)
        parser.add_argument("--show-phone", action="store_true")
        parser.add_argument("--dry-run", action="store_true")
        parser.add_argument("--report", help="Optional CSV report path.")
>>>>>>> 1d670fd (refactor)

    def handle(self, *args, **options):
        file_path = Path(options["file"])
        if not file_path.is_absolute():
            file_path = Path(settings.BASE_DIR) / file_path
<<<<<<< HEAD

=======
>>>>>>> 1d670fd (refactor)
        if not file_path.exists():
            raise CommandError(f"DOCX file not found: {file_path}")

        document = Document(file_path)
        if not document.tables:
            raise CommandError("The DOCX does not contain a member table.")

<<<<<<< HEAD
        table = document.tables[0]
        rows = []
        for row in table.rows[1:]:
            cells = [clean_text(cell.text) for cell in row.cells]
            if len(cells) < 5:
                continue

            serial, designation, name, address, phone = cells[:5]
            if not name:
                continue

            rows.append({
=======
        rows = []
        for row_number, row in enumerate(document.tables[0].rows[1:], start=2):
            cells = [clean_text(cell.text) for cell in row.cells]
            if len(cells) < 5:
                continue
            serial, designation, name, address, phone = cells[:5]
            if not name:
                continue
            rows.append({
                "row_number": row_number,
>>>>>>> 1d670fd (refactor)
                "serial": nepali_number_to_int(serial),
                "designation": designation,
                "name": name,
                "address": address,
<<<<<<< HEAD
                "phone": clean_phone(phone),
=======
                "phone": normalize_phone(phone),
>>>>>>> 1d670fd (refactor)
            })

        if not rows:
            raise CommandError("No committee members were found in the DOCX table.")

<<<<<<< HEAD
        unit_name = clean_text(options["unit_name"])
        membership_type = options["membership_type"]
        show_phone = options["show_phone"]

        created_count = 0
        updated_count = 0

        with transaction.atomic():
            if options["replace_unit"]:
                deleted_count, _ = Member.objects.filter(
                    level="rural_municipality",
                    unit_name=unit_name,
                ).delete()
                self.stdout.write(
                    self.style.WARNING(
                        f"Deleted {deleted_count} existing rows from {unit_name}."
                    )
                )

            for row in rows:
                member = Member.objects.filter(
                    level="rural_municipality",
                    unit_name=unit_name,
                    name_ne=row["name"],
                ).first()

                if member is None:
                    member = Member(
                        level="rural_municipality",
                        unit_name=unit_name,
                        name_ne=row["name"],
                        membership_type=membership_type,
                    )
                    created = True
                else:
                    # The DOCX does not state membership type, so keep the type
                    # already assigned to an existing registry record.
                    created = False

                member.name_en = member.name_en or romanize_nepali_name(row["name"])
                member.status = "active"
                member.municipality = unit_name
                member.address = row["address"]
                member.designation = row["designation"]
                member.phone = row["phone"]
                member.show_phone_publicly = show_phone
                member.is_public = True
                member.sort_order = row["serial"]
                member.save()

                if created:
                    created_count += 1
                else:
                    updated_count += 1

        self.stdout.write(
            self.style.SUCCESS(
                f"Imported {len(rows)} committee rows: "
                f"{created_count} created, {updated_count} updated."
            )
        )
        self.stdout.write(
            self.style.WARNING(
                "The DOCX does not provide membership type or membership numbers. "
                f"Type used: {membership_type}; new numbers were assigned inside "
                f"the {unit_name} registry."
            )
        )
        self.stdout.write(
            "Phone numbers were "
            + ("made public." if show_phone else "stored but hidden from the public directory.")
        )
=======
        results = []
        try:
            with transaction.atomic():
                unit = get_or_create_unit("rural_municipality", clean_text(options["unit_name"]))
                category = get_or_create_category(options["membership_type"])

                for data in rows:
                    source_key = hashlib.sha256(
                        f"{file_path.name}|{data['row_number']}|{data['name']}".encode("utf-8")
                    ).hexdigest()[:40]
                    existing = MembershipRecord.objects.select_related("person").filter(
                        organization_unit=unit,
                        source_identity_key=source_key,
                    ).first()

                    if existing:
                        person = existing.person
                        existing.designation = data["designation"]
                        existing.address_display = data["address"]
                        existing.is_public = True
                        existing.save()
                        status = "updated"
                    else:
                        identity = resolve_person_identity(
                            name_ne=data["name"],
                            phone=data["phone"],
                            location=data["address"],
                        )
                        person = identity.person
                        normalized = MembershipRecord.objects.create(
                            person=person,
                            category=category,
                            organization_unit=unit,
                            membership_number=str(data["serial"] or ""),
                            source_identity_key=source_key,
                            status="active",
                            designation=data["designation"],
                            address_display=data["address"],
                            is_public=True,
                        )
                        # Preserve/update the legacy table until all production code has migrated.
                        legacy = Member.objects.filter(
                            level="rural_municipality",
                            unit_name=unit.name_en,
                            membership_type=category.code,
                            membership_number=str(data["serial"] or ""),
                        ).first() or Member(
                            level="rural_municipality",
                            unit_name=unit.name_en,
                            membership_type=category.code,
                            membership_number=str(data["serial"] or ""),
                        )
                        legacy.name_ne = data["name"]
                        legacy.name_en = person.name_en
                        legacy.status = "active"
                        legacy.municipality = unit.name_en
                        legacy.address = data["address"]
                        legacy.designation = data["designation"]
                        legacy.phone = data["phone"]
                        legacy.show_phone_publicly = bool(options["show_phone"])
                        legacy.is_public = True
                        legacy.sort_order = data["serial"]
                        legacy.save()
                        normalized.legacy_member = legacy
                        normalized.save(update_fields=["legacy_member"])
                        status = "created"

                    results.append({**data, "status": status, "person_id": person.pk})

                if options["dry_run"]:
                    transaction.set_rollback(True)
        except Exception as exc:
            raise CommandError(str(exc)) from exc

        report_path = options.get("report")
        if report_path:
            path = Path(report_path)
            if not path.is_absolute():
                path = Path(settings.BASE_DIR) / path
            path.parent.mkdir(parents=True, exist_ok=True)
            with path.open("w", encoding="utf-8-sig", newline="") as destination:
                writer = csv.DictWriter(destination, fieldnames=list(results[0].keys()))
                writer.writeheader()
                writer.writerows(results)
            self.stdout.write(f"Report: {path}")

        prefix = "DRY RUN: " if options["dry_run"] else ""
        created = sum(row["status"] == "created" for row in results)
        updated = sum(row["status"] == "updated" for row in results)
        self.stdout.write(self.style.SUCCESS(
            f"{prefix}processed {len(results)} committee rows: {created} created, {updated} updated, 0 deleted."
        ))
        self.stdout.write(self.style.WARNING(
            "The source does not state membership category. Review the selected category and English spellings in admin."
        ))
>>>>>>> 1d670fd (refactor)
